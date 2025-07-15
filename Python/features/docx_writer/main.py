from docx import Document
import os
import sys

# Get the file path entered by the user
file_path = input("Enter source file path: ").strip()

if not file_path:
    print("Must enter source file path!")
    sys.exit(1)

# Check whether the file exists
if not os.path.exists(file_path):
    print(f"Error: File '{file_path}' does not exist.")
    sys.exit(1)

# Make sure it's a file, not a directory
if not os.path.isfile(file_path):
    print(f"Error: '{file_path}' is not a file.")
    sys.exit(1)

try:
    # Get the file name and extension
    file_dir, filename = os.path.split(file_path)
    filename_without_ext, ext = os.path.splitext(filename)

    # Read the contents of the file
    with open(file_path, "r", encoding="utf-8") as file:
        lines = file.readlines()

    # Create Word document
    doc = Document()

    # Reduce every lines
    for line in lines:
        line = line.strip()
        if line.startswith('#'):
            # Calculate the number of # and set the title level (1-6)
            level = min(line.count('#'), 6)
            title_text = line.lstrip('#').lstrip()
            doc.add_heading(title_text, level=level)
        else:
            doc.add_paragraph(line)

    # Ensure the existence of the output directory
    output_dir = os.path.join(os.getcwd(), 'docx')
    os.makedirs(output_dir, exist_ok=True)

    # Build the output file path
    output_filename = f'{filename_without_ext.replace(" ", "_")}.docx'
    output_path = os.path.join(output_dir, output_filename)

    # Save Word document
    doc.save(output_path)
    print(f"Successfully converted to: '{output_path}'")

except Exception as e:
    print(f"An error occurred: {e}")
    sys.exit(1)